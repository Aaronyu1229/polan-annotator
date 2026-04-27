"""產生 Amber 安裝步驟 .docx — 含可執行 cli 區塊。

執行：
    /tmp/docx-env/bin/python /tmp/make_amber_doc.py
輸出：
    ~/Desktop/Amber 安裝指引.docx
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade(paragraph, fill="F4F4F4"):
    """段落淺灰背景（cli code 區塊感）。"""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def code(doc, text):
    """加一個等寬字體 + 淺灰背景的 cli 區塊。"""
    p = doc.add_paragraph()
    shade(p)
    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def h(doc, text, level=1):
    """heading helper。"""
    doc.add_heading(text, level=level)


def step(doc, num, title, body=None):
    """『步驟 N — 標題』+ 描述。"""
    p = doc.add_paragraph()
    run = p.add_run(f"步驟 {num}　")
    run.bold = True
    run.font.size = Pt(14)
    run2 = p.add_run(title)
    run2.bold = True
    run2.font.size = Pt(14)
    if body:
        doc.add_paragraph(body)


doc = Document()

# 全域字體
style = doc.styles["Normal"]
style.font.name = "PingFang TC"
style.font.size = Pt(11)

# ── 封面 ──────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("珀瀾聲音標註工具")
r.bold = True
r.font.size = Pt(22)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("Amber Mac 本機安裝指引")
sr.font.size = Pt(14)
sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Aaron 在場操作　•　預估 30-45 分鐘　•　2026-04-26")
mr.font.size = Pt(10)
mr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ── 重要說明 ──────────────────────────
h(doc, "📌 開始之前", level=1)
doc.add_paragraph(
    "本文件兩種使用方式："
)
doc.add_paragraph(
    "方式 A（推薦）：跑「全自動安裝腳本」（步驟 7），90% 工作一次完成。",
    style="List Bullet",
)
doc.add_paragraph(
    "方式 B：照步驟 1-12 一個一個跑（需要自己 copy-paste cli）。",
    style="List Bullet",
)
doc.add_paragraph(
    "兩種都需要：（1）你帶來的 polan-data.tar.gz、（2）AirDrop 給她、（3）她讓你用她的 terminal。"
)
doc.add_paragraph()

doc.add_paragraph("✅ 你帶來的東西 checklist：").bold = True
doc.add_paragraph("• 你的 MacBook（已有 ~/Desktop/polan-data.tar.gz，406 MB）", style="List Bullet")
doc.add_paragraph("• 手機 LINE / 電話", style="List Bullet")
doc.add_paragraph("• 這份文件（已存她桌面或開 iCloud / Google Doc）", style="List Bullet")

doc.add_page_break()

# ── 步驟 1 ──
h(doc, "前置：在 Amber Mac 確認環境", level=1)
step(doc, 1, "確認 macOS 版本 + Xcode CLT")
doc.add_paragraph("打開「Terminal」應用程式（Cmd+空白 搜 terminal）→ 貼上這段：")
code(doc, "sw_vers && xcode-select -p")
doc.add_paragraph(
    "預期：macOS 11+ 都 OK；最後一行會印路徑。若印「xcode-select: error: ...」表示沒裝，需要先跑："
)
code(doc, "xcode-select --install")
doc.add_paragraph(
    "會跳 GUI 安裝視窗，按「安裝」→ 等 5-10 分鐘 → 裝完才繼續下一步。"
)

# ── 兩台 Mac 都打開 AirDrop ──
h(doc, "兩台 Mac 都開 AirDrop（先準備傳檔通道）", level=2)
step(doc, 2, "開 AirDrop")
doc.add_paragraph("兩台 Mac 都做：")
doc.add_paragraph("a. Finder → 選單列「前往」→ AirDrop（Cmd+Shift+R）", style="List Number")
doc.add_paragraph('b. 視窗最下方「讓我被以下對象發現」選 「所有人」', style="List Number")
doc.add_paragraph("c. 確認 WiFi + 藍牙都開（兩個都要）", style="List Number")

doc.add_page_break()

# ── 全自動 ──
h(doc, "🚀 方式 A：全自動安裝（推薦）", level=1)
doc.add_paragraph(
    "如果上面 Xcode CLT + AirDrop 都已 OK，只剩 4 步："
)

step(doc, 3, "你 Mac AirDrop 兩個檔給她")
doc.add_paragraph("把你桌面這兩個檔拖到她的 AirDrop 頭像上：")
doc.add_paragraph("• polan-data.tar.gz（406 MB，音檔 + DB）", style="List Bullet")
doc.add_paragraph("• install_amber_polan.sh（自動安裝腳本）", style="List Bullet")
doc.add_paragraph("她按「接受並儲存」 → 預設存 ~/Downloads")

step(doc, 4, "她 Mac terminal 跑自動腳本")
doc.add_paragraph("貼這段（會做 9 件事：裝 uv、clone repo、解壓、建 venv、裝套件、跑 tests、裝桌面 icon）：")
code(doc,
    "chmod +x ~/Downloads/install_amber_polan.sh && \\\n"
    "  ~/Downloads/install_amber_polan.sh"
)
doc.add_paragraph(
    "看到最後一行 ✅ ALL DONE 即成功。預估 5-10 分鐘（看網速 + 她 CPU）。"
)
doc.add_paragraph("⚠️ 中間若失敗，看 terminal 紅色錯誤訊息 → 跳到下面「方式 B」對應步驟手動跑。")

step(doc, 5, "雙擊桌面 icon 啟動")
doc.add_paragraph(
    "桌面有「啟動珀瀾標註工具.command」icon → 第一次雙擊會跳「無法驗證開發者」 → "
    "系統設定 → 隱私權與安全性 → 捲到底「強制打開」 → 再雙擊 → 黑視窗 + 瀏覽器自動開"
)

step(doc, 6, "驗收 4 件事（瀏覽器裡）")
doc.add_paragraph("a. 清單顯示「已標 11 / 33」", style="List Number")
doc.add_paragraph("b. 點任一她已標的音檔 → 音源類型標題顯示「（可多選，至少選一項）」+ chip 已勾 ambience", style="List Number")
doc.add_paragraph("c. Loop Capability 顯示為 checkbox（不是 radio）+ 已勾上次選的", style="List Number")
doc.add_paragraph("d. 遊戲類型 (genre) 是 chip 區塊（chip + 加入按鈕，不是純文字輸入）", style="List Number")

doc.add_page_break()

# ── 方式 B 手動 ──
h(doc, "🛠 方式 B：手動逐步（自動腳本失敗時備案）", level=1)

step(doc, 7, "裝 uv（Python 套件管理）")
code(doc,
    "curl -LsSf https://astral.sh/uv/install.sh | sh\n"
    "source ~/.zshrc\n"
    "uv --version"
)

step(doc, 8, "Clone repo")
code(doc,
    "cd ~/Desktop\n"
    "git clone https://github.com/Aaronyu1229/polan-annotator.git\n"
    "cd polan-annotator\n"
    "git log --oneline -1   # 應看到 00e62b5 [Phase 5 #5]..."
)

step(doc, 9, "解壓 AirDrop 傳來的資料")
code(doc,
    "cd ~/Desktop/polan-annotator\n"
    "tar -xzf ~/Downloads/polan-data.tar.gz\n"
    "ls data/audio/ | wc -l                                           # 33\n"
    'sqlite3 data/annotations.db "SELECT COUNT(*) FROM annotation;"   # 11'
)

step(doc, 10, "建 venv + 裝依賴 + 跑測試")
code(doc,
    "uv venv\n"
    "source .venv/bin/activate\n"
    "uv pip install -e .                  # 等 1-3 分鐘\n"
    "pytest -q 2>&1 | tail -3             # 應 84 passed"
)

step(doc, 11, "裝桌面 icon")
code(doc,
    "bash scripts/install_desktop_shortcut.sh"
)
doc.add_paragraph(
    "桌面會多「啟動珀瀾標註工具.command」 → 雙擊 → "
    "第一次跳「無法驗證」 → 系統設定 → 隱私權與安全性 → 強制打開 → 再雙擊"
)

doc.add_page_break()

# ── 你離開前 ──
h(doc, "✅ 你離開前 checklist", level=1)
doc.add_paragraph("□ 雙擊桌面 icon 順 1 次給她看", style="List Bullet")
doc.add_paragraph("□ 讓她「自己」雙擊一次（你站旁邊）", style="List Bullet")
doc.add_paragraph("□ 她開一個沒標過的音檔，全程標完按「儲存並下一個」成功", style="List Bullet")
doc.add_paragraph(
    "□ DB 確認 +1：sqlite3 ~/Desktop/polan-annotator/data/annotations.db "
    '"SELECT COUNT(*) FROM annotation;" → 12',
    style="List Bullet",
)
doc.add_paragraph("□ 你 LINE / 手機備好給她（卡住找你）", style="List Bullet")

# ── 教 Amber ──
h(doc, "📋 教 Amber（便利貼黏螢幕邊）", level=1)
table = doc.add_table(rows=2, cols=2)
table.style = "Light Grid Accent 1"
table.cell(0, 0).text = "動作"
table.cell(0, 1).text = "怎麼做"
table.cell(1, 0).text = "開工"
table.cell(1, 1).text = "雙擊桌面「啟動珀瀾標註工具」icon"
hdr = table.rows[0].cells
for c in hdr:
    for p in c.paragraphs:
        for run in p.runs:
            run.bold = True
table.add_row()
table.rows[2].cells[0].text = "結束"
table.rows[2].cells[1].text = "終端機視窗按 Ctrl+C → 關視窗"

# ── 你回家後 ──
h(doc, "🏠 你回家後（在你自己 Mac）", level=1)
doc.add_paragraph("關掉之前在你 Mac 跑的舊 server + cloudflared tunnel：")
code(doc,
    "pkill -f cloudflared\n"
    'pkill -f "uvicorn src.main:app"'
)
doc.add_paragraph("LINE 她：")
note = doc.add_paragraph()
shade(note, fill="FFF8DC")
nr = note.add_run(
    "「桌面 icon 從現在起是你的工作入口，那個 Cloudflare 連結之後不會再開了，不要再用喔。」"
)
nr.italic = True

# ── 踩坑表 ──
h(doc, "⚠️ 可能踩到的坑", level=1)
table2 = doc.add_table(rows=1, cols=2)
table2.style = "Light Grid Accent 2"
table2.cell(0, 0).text = "狀況"
table2.cell(0, 1).text = "解法"
hdr = table2.rows[0].cells
for c in hdr:
    for p in c.paragraphs:
        for run in p.runs:
            run.bold = True

troubles = [
    ("uv pip install 卡很久", "等。第一次裝 numpy/scipy 編譯 2-3 分鐘很正常"),
    ("pytest audio_analysis fail", "data/audio 沒解開或 path 錯 → 重看步驟 9"),
    ("雙擊 icon 黑視窗閃過就關", "路徑問題 → 重跑「bash scripts/install_desktop_shortcut.sh」"),
    ("瀏覽器沒自動開", "手動打 http://localhost:8000/?annotator=amber"),
    ("她說「介面跟舊的一樣」", "瀏覽器按 Cmd+Shift+R 強制重整"),
    ("AirDrop 找不到對方", "兩台距離拉近、確認「所有人」可見、WiFi+藍牙都開"),
    ("xcode-select 安裝視窗沒跳", "App Store 搜「Xcode Command Line Tools」手動裝"),
]
for s, r in troubles:
    row = table2.add_row().cells
    row[0].text = s
    row[1].text = r

OUT = Path.home() / "Desktop" / "Amber 安裝指引.docx"
doc.save(OUT)
print(f"✅ 寫入：{OUT}")
